import time
from docx import Document  # 引入docx类生成docx文档
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

__version__ = "1.0.0"


class PrintPreview:
    """本类负责生成完整的口算题文档使之适合打印机打印。可以生成多套题，生成数可以控。

    - @p_list   list
    需要打印口算题库，至少包含一套口算题

    - @p_title   list
    页面标题，这个标题的生成依据程序题型的选择和数字的范围选择而生成，例如：选择了0-20，加减法，进退位
    则自动生成标题为：0到20加减法进退位混合口算题，list中包含了多套题的页面标题名称

    - @p_column  int
    打印页排版口算题的列数

    """

    p_list = None
    p_title = None
    p_subtitle = None
    p_column = None
    p_title_size = None
    p_subtitle_size = None
    p_content_siae = None

    def __init__(self, l, tit, subtitle, out_put, col=3, tsize=26, subsize=11, csize=16):
        """
        :param l: list 需要打印的口算题列表
        :param tit: list 口算页标题
        :param subtitle str 小标题
        :param out_put str 保存目录
        :param col: int 列数
        :param tsize: int 标题字号
        :param csize: int 口算题字号
        """
        self.p_list = l
        self.p_title = tit
        self.p_subtitle = subtitle
        self.p_out_put = Path(out_put)
        self.p_column = col
        self.p_title_size = tsize
        self.p_subtitle_size = subsize
        self.p_content_siae = csize

    def create_psmdocx(self, l, title, docxname):
        """
        :param l list 一组题库
        :param title str 页面标题
        :param docxname  str 题库保存文件名
        :return: none
        """
        if title == '':
            page_title = '小学生口算题'
        else:
            page_title = title
        p_docx = Document()  # 创建一个docx文档
        p_docx.styles['Normal'].font.name = u'Times'  # 可换成word里面任意字体
        p = p_docx.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
        run = p.add_run(page_title)
        run.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        run.font.size = Pt(self.p_title_size)  # 字体大小设置，和word里面的字号相对应

        sp = p_docx.add_paragraph()
        sp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
        srun = sp.add_run(self.p_subtitle)
        srun.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        srun.font.size = Pt(self.p_subtitle_size)  # 字体大小设置，和word里面的字号相对应

        # 判断需要用到的行数
        if len(l) % self.p_column:
            rs = len(l) // self.p_column + 2
        else:
            rs = len(l) // self.p_column + 1

        # print(rs)

        # 将口算题添加到docx表格中
        k = 0  # 计数器
        table = p_docx.add_table(rows=rs, cols=self.p_column)

        for i in range(rs):
            if i > 0:
                row_cells = table.rows[i].cells
                for j in range(self.p_column):
                    if k > len(l) - 1:
                        break
                    else:
                        row_cells[j].text = l[k]
                        k = k + 1
        table.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style.font.color.rgb = RGBColor(54, 0, 0)  # 颜色设置，这里是用RGB颜色
        table.style.font.size = Pt(self.p_content_siae)  # 字体大小设置，和word里面的字号相对应

        save_file = self.p_out_put / ("%s.docx" % docxname)
        p_docx.save(save_file)  # 输出docx

    def produce(self):
        for p in self.p_list:
            t = time.strftime("%y%m%d%H%M%S", time.localtime())
            self.create_psmdocx(p, self.p_title, "%s_%s" % (self.p_title, t))

